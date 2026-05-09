from typing import List, Dict, Any, Optional
import json
import logging
from sentence_transformers import SentenceTransformer
import chromadb
from chromadb.config import Settings
import os
from dotenv import load_dotenv
import docx   

# Load environment variables
load_dotenv()

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)
# model = SentenceTransformer('all-MiniLM-L6-v2')  <-- REMOVED GLOBAL INIT
model = None

def get_model():
    global model
    if not model:
        # Force offline mode to prevent connection timeouts
        os.environ["HF_HUB_OFFLINE"] = "1"
        os.environ["TRANSFORMERS_OFFLINE"] = "1"
        os.environ["HF_DATASETS_OFFLINE"] = "1"
        logger.info("⏳ Loading SentenceTransformer model (STRICT OFFLINE MODE)...")
        try:
            model = SentenceTransformer('all-MiniLM-L6-v2')
            logger.info("✅ Model loaded.")
        except Exception as e:
            logger.warning(f"Offline load failed: {e}. Retrying online...")
            # Unset variables to allow online fallback
            os.environ.pop("HF_HUB_OFFLINE", None)
            os.environ.pop("TRANSFORMERS_OFFLINE", None)
            os.environ.pop("HF_DATASETS_OFFLINE", None)
            model = SentenceTransformer('all-MiniLM-L6-v2')
            logger.info("✅ Model loaded (Online fallback).")
    return model

# Use absolute path relative to this file to ensure it works from ANY working directory

current_dir = os.path.dirname(os.path.abspath(__file__))
CHROMA_DB_PATH = os.path.join(current_dir, "chroma_db")

OBJECTS_COLLECTION = "salesforce_objects"
FIELDS_COLLECTION = "salesforce_fields"
 
chroma_client = chromadb.PersistentClient(
    path=CHROMA_DB_PATH,
    settings=Settings(
        anonymized_telemetry=False,
        allow_reset=True
    )
)

 
schema_data = None


# ========================================
# HELPER FUNCTIONS (DEFINE FIRST!)
# ========================================

def extract_json_from_word(file_path):
    """Extract JSON from Word document with error handling"""
    if not file_path:
        raise ValueError("File path cannot be empty")
        
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"Word document not found: {file_path}")
        
    try:
        doc = docx.Document(file_path)
        text = ""
        for para in doc.paragraphs:
            text += para.text + "\n"
        
        # Also extract from tables if present
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + "\n"
        
        if not text.strip():
            raise ValueError("No text content found in Word document")
            
        logger.info(f"Successfully extracted text from {file_path}")
        return text
        
    except Exception as e:
        logger.error(f"Error extracting JSON from Word document {file_path}: {e}")
        raise RuntimeError(f"Failed to extract content from Word document: {e}")


def parse_json_from_text(text):
    """Parse JSON from extracted text with error handling"""
    if not text or not text.strip():
        logger.warning("No text provided for JSON parsing")
        return None
        
    try:
        json_start = text.find('[')
        json_end = text.rfind(']') + 1
        
        if json_start == -1 or json_end == 0:
            logger.warning("No JSON array found in text")
            return None
            
        json_data = text[json_start:json_end]
        
        if not json_data.strip():
            logger.warning("Empty JSON data extracted")
            return None
            
        parsed_data = json.loads(json_data)
        logger.info(f"Successfully parsed JSON with {len(parsed_data)} items")
        return parsed_data
        
    except json.JSONDecodeError as e:
        logger.error(f"Invalid JSON format: {e}")
        return None
    except Exception as e:
        logger.error(f"Error parsing JSON from text: {e}")
        return None


# ========================================
# CHROMADB MANAGER CLASS
# ========================================

class ChromaDBManager:
    """Centralized ChromaDB management with best practices and error handling"""

    def __init__(self, chroma_client):
        self.client = chroma_client
        self.objects_collection = None
        self.fields_collections = {}

    def get_or_create_objects_collection(self):
        """Get or create objects collection with proper error handling"""
        if not self.objects_collection:
            try:
                self.objects_collection = self.client.get_collection(
                    name=OBJECTS_COLLECTION,
                    embedding_function=None
                )
                logger.info(f"Retrieved existing objects collection: {OBJECTS_COLLECTION}")
            except Exception as get_error:
                try:
                    self.objects_collection = self.client.create_collection(
                        name=OBJECTS_COLLECTION,
                        embedding_function=None,
                        metadata={"description": "Salesforce object embeddings"}
                    )
                    logger.info(f"Created new objects collection: {OBJECTS_COLLECTION}")
                except Exception as create_error:
                    logger.error(f"Failed to create objects collection: {create_error}")
                    raise RuntimeError(f"Cannot create objects collection: {create_error}")
        return self.objects_collection

    def get_or_create_fields_collection(self, object_name: str):
        """Get or create fields collection for specific object with error handling"""
        if not object_name:
            raise ValueError("Object name cannot be empty")
            
        collection_name = f"{FIELDS_COLLECTION}_{object_name.lower()}"
        
        if object_name not in self.fields_collections:
            try:
                self.fields_collections[object_name] = self.client.get_collection(
                    name=collection_name,
                    embedding_function=None
                )
                logger.info(f"Retrieved existing fields collection: {collection_name}")
            except Exception as get_error:
                try:
                    self.fields_collections[object_name] = self.client.create_collection(
                        name=collection_name,
                        embedding_function=None,
                        metadata={
                            "description": f"Fields for {object_name} object",
                            "object_name": object_name
                        }
                    )
                    logger.info(f"Created new fields collection: {collection_name}")
                except Exception as create_error:
                    logger.error(f"Failed to create fields collection for {object_name}: {create_error}")
                    raise RuntimeError(f"Cannot create fields collection for {object_name}: {create_error}")
        
        return self.fields_collections[object_name]

    def get_existing_ids(self, collection, ids: List[str]) -> set:
        """Get existing IDs from collection with error handling"""
        existing = set()
        if not ids:
            return existing
            
        try:
            results = collection.get(ids=ids, include=[])
            if results and 'ids' in results:
                existing.update(results['ids'])
        except Exception as e:
            logger.warning(f"Error checking existing IDs: {e}")
            # Return empty set to allow processing to continue
        return existing

    def store_object_embeddings(self, schema_data):
        """Store object embeddings with comprehensive error handling"""
        if not schema_data:
            logger.warning("No schema data provided for embedding storage")
            return
            
        try:
            collection = self.get_or_create_objects_collection()
            objects_to_add = []
            embeddings_to_add = []
            metadatas_to_add = []
            ids_to_add = []

            temp_ids = [f"obj_{idx}_{item['object']}" for idx, item in enumerate(schema_data)]
            existing_ids = self.get_existing_ids(collection, temp_ids)

            for idx, item in enumerate(schema_data):
                try:
                    object_name = item.get('object')
                    if not object_name:
                        logger.warning(f"Skipping item {idx}: missing object name")
                        continue
                        
                    doc_id = f"obj_{idx}_{object_name}"

                    if doc_id in existing_ids:
                        continue

                    embedding = get_model().encode(object_name).tolist()
                    
                    objects_to_add.append(object_name)
                    embeddings_to_add.append(embedding)
                    metadatas_to_add.append({
                        "object_name": object_name,
                        "field_count": len(item.get('fields', [])),
                        "schema_index": idx
                    })
                    ids_to_add.append(doc_id)
                    
                except Exception as e:
                    logger.warning(f"Error processing object {item.get('object', 'unknown')} at index {idx}: {e}")
                    continue

            if objects_to_add:
                try:
                    batch_size = 100
                    for i in range(0, len(objects_to_add), batch_size):
                        collection.add(
                            documents=objects_to_add[i:i + batch_size],
                            embeddings=embeddings_to_add[i:i + batch_size],
                            metadatas=metadatas_to_add[i:i + batch_size],
                            ids=ids_to_add[i:i + batch_size]
                        )
                    logger.info(f"Successfully stored {len(objects_to_add)} new object embeddings")
                except Exception as e:
                    logger.error(f"Error storing embeddings in batches: {e}")
                    raise RuntimeError(f"Failed to store object embeddings: {e}")
            else:
                logger.info("No new object embeddings to store")
                
        except Exception as e:
            logger.error(f"Error in store_object_embeddings: {e}")
            raise


    def store_field_embeddings(self, schema_data):
        """Store field embeddings with comprehensive error handling"""
        if not schema_data:
            logger.warning("No schema data provided for field embedding storage")
            return
            
        processed_objects = 0
        failed_objects = []
        
        for item in schema_data:
            try:
                object_name = item.get('object')
                if not object_name:
                    logger.warning("Skipping item: missing object name")
                    continue
                    
                fields = item.get('fields', [])
                if not fields:
                    logger.info(f"No fields found for object: {object_name}")
                    continue

                collection = self.get_or_create_fields_collection(object_name)
                temp_ids = [f"field_{object_name}_{idx}_{field.get('apiname', 'unknown')}" 
                           for idx, field in enumerate(fields) if field.get('apiname')]
                # FORCE UPDATE REVERTED: Re-enabling caching to fix startup crash
                existing_ids = self.get_existing_ids(collection, temp_ids)
                # existing_ids = set() 


                documents_to_add = []
                embeddings_to_add = []
                metadatas_to_add = []
                ids_to_add = []

                for idx, field in enumerate(fields):
                    try:
                        field_name = field.get('apiname')
                        if not field_name:
                            logger.warning(f"Skipping field {idx} in {object_name}: missing apiname")
                            continue
                            
                        doc_id = f"field_{object_name}_{idx}_{field_name}"
                        if doc_id in existing_ids:
                            continue

                        description = field.get('description', '')
                        datatype = field.get('datatype', '')
                        default_value = field.get('defaultValue', '')
                        label = field.get('FieldLabel', field_name) # Uses enriched label or fallback to name
                        
                        # Extract needvalue / isrequired
                        # Checking various casing possibilities including sloppy whitespace
                        need_value = (field.get('needvalue') or 
                                     field.get('needvalue ') or 
                                     field.get(' needvalue ') or 
                                     field.get('needValue') or 
                                     field.get('NeedValue')) or False
                        # Convert to boolean
                        if isinstance(need_value, str):
                            need_value = need_value.lower() == 'true'
                        
                        combined_text = f"fieldapiname: {field_name}, label: {label}, description: {description}, datatype: {datatype}, needvalue: {need_value}"
                        
                        embedding = get_model().encode(combined_text).tolist()

                        documents_to_add.append(combined_text)
                        embeddings_to_add.append(embedding)
                        metadatas_to_add.append({
                            "field_name": field_name,
                            "FieldLabel": label, # Persist FieldLabel
                            "object_name": object_name,
                            "description": description,
                            "datatype": datatype,
                            "defaultValue": default_value,
                            "field_index": idx,
                            "needvalue": bool(need_value)
                        })
                        ids_to_add.append(doc_id)
                        
                    except Exception as e:
                        logger.warning(f"Error processing field {field.get('apiname', 'unknown')} in {object_name}: {e}")
                        continue

                if documents_to_add:
                    try:
                        batch_size = 100
                        for i in range(0, len(documents_to_add), batch_size):
                            collection.upsert( # Use upsert to update metadata
                                documents=documents_to_add[i:i+batch_size],
                                embeddings=embeddings_to_add[i:i+batch_size],
                                metadatas=metadatas_to_add[i:i+batch_size],
                                ids=ids_to_add[i:i+batch_size]
                            )
                        logger.info(f"Stored {len(documents_to_add)} new/updated field embeddings for {object_name}")
                        processed_objects += 1
                    except Exception as e:
                        logger.error(f"Error storing field embeddings for {object_name}: {e}")
                        failed_objects.append(object_name)
                else:
                    logger.info(f"No new field embeddings to store for {object_name}")
                    
            except Exception as e:
                logger.error(f"Error processing object {item.get('object', 'unknown')}: {e}")
                failed_objects.append(item.get('object', 'unknown'))
                continue
        
        logger.info(f"Field embedding storage complete. Processed: {processed_objects}, Failed: {len(failed_objects)}")
        if failed_objects:
            logger.warning(f"Failed objects: {failed_objects}")

    def get_need_value_fields(self, object_name: str) -> List[Dict]:
        """Get fields marked as needvalue=True for a specific object"""
        if not object_name:
            return []
            
        try:
            collection = self.get_or_create_fields_collection(object_name)
            # Fetch all fields with needvalue=True
            results = collection.get(
                where={"needvalue": True},
                include=["metadatas"]
            )
            
            if not results or not results.get("metadatas"):
                return []
                
            return [m for m in results["metadatas"]]
            
        except Exception as e:
            logger.error(f"Error fetching needvalue fields for {object_name}: {e}")
            return []


    def search_objects(self, query: str, top_k: int = 2) -> List[Dict]:
        """Search objects with error handling"""
        if not query or not query.strip():
            logger.warning("Empty query provided for object search")
            return []
            
        try:
            collection = self.get_or_create_objects_collection()
            query_embedding = get_model().encode(query).tolist()
            
            results = collection.query(
                query_embeddings=[query_embedding],
                n_results=top_k,
                include=["documents", "metadatas", "distances"]
            )
            
            if not results.get('documents') or not results['documents'][0]:
                logger.info(f"No objects found for query: {query}")
                return []
            
            formatted_results = []
            for i in range(len(results['documents'][0])):
                try:
                    formatted_results.append({
                        "object_name": results['metadatas'][0][i]['object_name'],
                        "distance": results['distances'][0][i],
                        "metadata": results['metadatas'][0][i]
                    })
                except (KeyError, IndexError) as e:
                    logger.warning(f"Error formatting result {i}: {e}")
                    continue
                    
            return formatted_results
            
        except Exception as e:
            logger.error(f"Error searching objects with query '{query}': {e}")
            return []

    def search_fields(self, object_name: str, query: str, top_k: int = 5) -> List[Dict]:
        """Search fields with comprehensive error handling"""
        if not object_name or not object_name.strip():
            logger.warning("Empty object name provided for field search")
            return []
            
        if not query or not query.strip():
            logger.warning("Empty query provided for field search")
            return []
            
        try:
            collection = self.get_or_create_fields_collection(object_name)
            query_embedding = get_model().encode(query).tolist()
            
            results = collection.query(
                query_embeddings=[query_embedding],
                n_results=top_k,
                include=["documents", "metadatas", "distances"]
            )
            
            if not results.get('documents') or not results['documents'][0]:
                logger.info(f"No fields found for object '{object_name}' with query '{query}'")
                return []
            
            formatted_results = []
            for i in range(len(results['documents'][0])):
                try:
                    metadata = results['metadatas'][0][i]
                    formatted_results.append({
                        "field_name": metadata.get('field_name', 'unknown'),
                        "FieldLabel": metadata.get('FieldLabel', metadata.get('field_name', 'unknown')),
                        "object_name": metadata.get('object_name', object_name),
                        "description": metadata.get('description', ''),
                        "datatype": metadata.get('datatype', ''),
                        "distance": results['distances'][0][i],
                        "document": results['documents'][0][i]
                    })
                except (KeyError, IndexError) as e:
                    logger.warning(f"Error formatting field result {i}: {e}")
                    continue
                    
            return formatted_results
            
        except Exception as e:
            logger.error(f"Error searching fields for {object_name} with query '{query}': {e}")
            return []

    def reset_collections(self):
        """Reset all collections with error handling"""
        try:
            self.client.reset()
            logger.info("All ChromaDB collections reset successfully")
            # Clear cached collections
            self.objects_collection = None
            self.fields_collections = {}
        except Exception as e:
            logger.error(f"Error resetting collections: {e}")
            raise RuntimeError(f"Failed to reset collections: {e}")


# Initialize ChromaDB Manager with error handling
try:
    chroma_manager = ChromaDBManager(chroma_client)
    logger.info("ChromaDB Manager initialized successfully")
except Exception as e:
    logger.critical(f"Failed to initialize ChromaDB Manager: {e}")
    raise SystemExit(f"ChromaDB Manager initialization failed: {e}")


# ========================================
# SCHEMA INITIALIZATION FUNCTION
# ========================================
 

def initialize_schema(force=False):
    """Initialize schema embeddings in ChromaDB with comprehensive error handling"""
    global schema_data
    try:
        # Go up 3 levels from .../mcp_module/Salesforcemcp/chromadbutils.py to .../Marketing agent/
        ROOT_DIR = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
        SCHEMA_PATH = os.path.join(ROOT_DIR, "schema_metadata.json")
        
        # Extract Salesforce schema JSON 
        logger.info(f"Loading salesforce schema from JSON file: {SCHEMA_PATH}")
        try:
            # Example of loading the schema
            with open(SCHEMA_PATH, "r") as f:
                schema_data = json.load(f)
            if not schema_data:
                logger.error("Schema file is empty or invalid")
                return False
        except FileNotFoundError:
            logger.error("Schema file not found: salesforce_schema.json")
            return False
        except PermissionError:
            logger.error("Permission denied accessing schema file: salesforce_schema.json")
            return False
        except json.JSONDecodeError as e:
            logger.error(f"Invalid JSON format in schema file: {e}")
            return False
        except Exception as e:
            logger.error(f"Unexpected error loading schema file: {e}")
            return False
        logger.info(f"Loaded salesforce schema with {len(schema_data)} top-level objects")
        
        # Store embeddings in ChromaDB with error handling
        logger.info("Storing object embeddings in ChromaDB...")
        try:
            chroma_manager.store_object_embeddings(schema_data)
        except Exception as e:
            logger.error(f"Error storing object embeddings: {e}")
            return False
        
        logger.info("Storing field embeddings in ChromaDB...")
        try:
            chroma_manager.store_field_embeddings(schema_data)
        except Exception as e:
            logger.error(f"Error storing field embeddings: {e}")
            return False
        
        if not schema_data:
            logger.warning("No schema data available to export")
        else:
            try:
                # 🏷️ ENRICHMENT: Generate/Ensure Labels exist before saving
                schema_data = enrich_schema_with_labels(schema_data)
                logger.info("✅ Schema enriched with labels")
            except Exception as e:
                logger.warning(f"Could not export schema metadata JSON: {e}")
                
        logger.info("Schema embeddings initialized successfully in ChromaDB")
        return True
        
    except Exception as e:
        logger.error(f"Unexpected error initializing schema: {e}")
        return False

# ========================================
# LABEL ENRICHMENT HELPER
# ========================================

def enrich_schema_with_labels(schema_data):
    """
    Ensures every field has a 'FieldLabel'.
    If 'FieldLabel' is missing, it auto-generates it from 'apiname'.
    Standardizes on 'FieldLabel' key and removes 'label'.
    """
    import re
    if not schema_data: return []
    
    for obj in schema_data:
        fields = obj.get('fields', [])
        for f in fields:
            # 1. Check for user-provided label keys (prioritize FieldLabel)
            user_label = (f.get('FieldLabel') or 
                          f.get('fieldLabel') or 
                          f.get('fieldlabel') or 
                          f.get('Label') or 
                          f.get('label'))
            
            # 2. Set standardised FieldLabel
            if user_label and str(user_label).strip():
                f['FieldLabel'] = str(user_label).strip()
            
            # 3. If NO label found, auto-generate from APIName
            if not f.get('FieldLabel'):
                apiname = f.get('apiname', '')
                if apiname:
                    # Remove __c suffix
                    clean_name = apiname.replace('__c', '')
                    # Split CamelCase
                    label_text = re.sub(r'((?<=[a-z])[A-Z]|(?<!\A)[A-Z](?=[a-z]))', r' \1', clean_name)
                    f['FieldLabel'] = label_text.strip()
                else:
                    f['FieldLabel'] = "Unknown Field"
            
            # 4. Remove duplicate 'label' key if it exists to keep JSON clean
            if 'label' in f:
                del f['label']
                    
    logger.info("✅ Schema enriched with FieldLabel (standardized)")
    return schema_data




# ========================================
# LAZY INITIALIZATION HELPER
# ========================================

_is_initialized = False

def ensure_schema_initialized():
    """
    Public method to ensure schema is loaded ONLY when needed.
    This prevents heavy ChromaDB startup during module import.
    """
    global _is_initialized
    if _is_initialized:
        return True
        
    logger.info("💤 Lazy-loading Triggered: Initializing Schema now...")
    try:
        if initialize_schema(force=False):
            _is_initialized = True
            return True
    except Exception as e:
        logger.error(f"Lazy initialization failed: {e}")
        
    return False


# ========================================
# INITIALIZE SCHEMA ON STARTUP
# ========================================

if __name__ == "__main__":
    # If run as a script, FORCE a clean rebuild
    logger.info("🔧 Manual Execution detected: Performing CLEAN REBUILD of Database...")
    try:
        chroma_manager.reset_collections()
        logger.info("🧹 Collections reset successfully.")
        
        # Force rebuild
        if initialize_schema(force=True):
             logger.info("✅ Database rebuilt and schema_metadata.json exported!")
        else:
             logger.error("❌ Failed to rebuild database.")
    except Exception as e:
        logger.error(f"Error during manual rebuild: {e}")

# NOTE: The 'else' block for automatic import initialization has been REMOVED.
# We now rely on 'ensure_schema_initialized()' being called by tools that need it.

